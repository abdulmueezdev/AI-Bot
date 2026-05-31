import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Styling Utility ────────────────────────────────────────────────────────
def apply_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for h_level in range(1, 4):
        h_style = doc.styles[f'Heading {h_level}']
        h_font = h_style.font
        h_font.name = 'Calibri'
        h_font.color.rgb = RGBColor(139, 0, 0) # Dark Crimson #8B0000

# ── Document 1: Handoff ─────────────────────────────────────────────────────
def create_handoff_doc(output_path):
    doc = Document()
    apply_styles(doc)

    # Cover Page
    title = doc.add_heading('DIGITAL CLONE AI — COMPLETE PROJECT HANDOFF', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\nAuthor: Antigravity\nDate: May 2026\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. What the project is
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph('The Digital Clone AI is an advanced, multi-tenant RAG (Retrieval-Augmented Generation) chatbot system. It serves as a digital resurrection of historical or fictional personas—currently tuned specifically for Franz Kafka (Alucard). It leverages a 3-tier memory system, strict token budgeting, and few-shot prompt injection to enforce absolute persona fidelity.')

    # 2. Phase Breakdown
    doc.add_heading('2. Phase Breakdown', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('Phase 1: GSD (Get S**t Done) - Backend Foundation\n').bold = True
    p1.add_run('Built the core FastAPI backend, Supabase pgvector integration, Gemini embeddings, and OpenRouter/Groq LLM streaming. Implemented the base ingestion pipeline and standard API contracts.')
    p2 = doc.add_paragraph()
    p2.add_run('Phase 2 & 3: Safety, Orchestration, and Calendar\n').bold = True
    p2.add_run('Added strict safety checks (Prompt Injection detection, Toxicity). Integrated Google Calendar sync capabilities and a robust orchestrator logic block to handle complex tool selections.')
    p3 = doc.add_paragraph()
    p3.add_run('Phase 4A-4C: Kafka Persona Tuning\n').bold = True
    p3.add_run('Fixed major chunking bugs that were truncating the text corpus. Reingested 4,489 chunks into Supabase. Rebuilt the prompt builder to strictly enforce a 3600 token budget. Injected 15 expert-curated few-shot conversation examples directly into the message roles to force Llama 3 to mimic Kafka’s brief, private, and melancholic letter-writing style.')
    p4 = doc.add_paragraph()
    p4.add_run('Phase 4D: UI/UX Overhaul\n').bold = True
    p4.add_run('Completely rewrote the Next.js frontend (ChatWindow.tsx, layout.tsx, MessageBubble.tsx) to match a v0.dev-inspired dark theme. Replaced default fonts with EB Garamond, added a pure black (#0a0a0a) background, and added auto-resizing shadcn text areas.')

    # 3. File Structure
    doc.add_heading('3. File Structure', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Path'
    hdr_cells[1].text = 'Purpose'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(139, 0, 0)
                run.font.bold = True

    files = [
        ('backend/app/main.py', 'FastAPI entry point, mounts routers and middleware.'),
        ('backend/app/llm_client.py', 'Groq/OpenRouter streaming logic and few-shot injection.'),
        ('backend/app/prompt_builder.py', 'Strict token budgeting and prompt assembly logic.'),
        ('backend/app/vector_store.py', 'Supabase pgvector CRUD operations.'),
        ('backend/app/ingest.py', 'Chunking and embedding pipeline using Gemini.'),
        ('backend/clones/alucard/config.yaml', 'Version 3.0 persona definition and few-shot examples.'),
        ('frontend/app/layout.tsx', 'Next.js Root layout, injects EB Garamond font.'),
        ('frontend/app/page.tsx', 'Main entry point mounting the ChatWindow.'),
        ('frontend/components/ChatWindow.tsx', 'Redesigned v0-style dark theme UI for interaction.'),
        ('frontend/lib/api.ts', 'Frontend API wrapper to communicate with backend.'),
    ]
    for path, purpose in files:
        row_cells = table.add_row().cells
        row_cells[0].text = path
        row_cells[1].text = purpose

    # 4. 3-Tier Memory System
    doc.add_heading('4. 3-Tier Memory System', level=1)
    doc.add_paragraph('1. Core Identity (YAML Config): Static rules, system instructions, and few-shot examples defining the persona bounds.')
    doc.add_paragraph('2. Episodic/Knowledge Memory (Supabase pgvector): The historical corpus. Messages are embedded by Gemini and semantic similarity retrieves the top 5 relevant chunks (Max 1200 tokens).')
    doc.add_paragraph('3. Semantic/Short-term History (Session): The recent chat context window, truncated oldest-first when exceeding the 600-token budget.')

    # 5. Token Budget Breakdown
    doc.add_heading('5. Token Budget Breakdown', level=1)
    doc.add_paragraph('Total Absolute Limit: 3600 Tokens (cl100k_base encoding)')
    t_table = doc.add_table(rows=1, cols=2)
    t_table.style = 'Table Grid'
    t_hdr = t_table.rows[0].cells
    t_hdr[0].text = 'Block'
    t_hdr[1].text = 'Budget'
    for cell in t_hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(139, 0, 0)
                run.font.bold = True
    budgets = [('Identity', '300'), ('Few-Shot', '800'), ('Calendar', '300'), ('Entity', '200'), ('Memory', '400'), ('Knowledge (RAG)', '1200'), ('History', '600'), ('Query', '500')]
    for b, v in budgets:
        row = t_table.add_row().cells
        row[0].text = b
        row[1].text = v

    # 6. Corpus
    doc.add_heading('6. Kafka Corpus', level=1)
    doc.add_paragraph('Total chunks ingested: 4,489')
    c_list = ['The_Metamorphosis.txt', 'The_trail.txt', 'Letter_to_my_father.txt', 'Daires_of_franz_kafka_1910-1913.txt', 'Franz_kafka_letter_to_felica.txt', 'complete_short_stories.txt', 'frans_kafka_milenaya_mektublar-eng.txt']
    for c in c_list:
        doc.add_paragraph(c, style='List Bullet')

    # 7. Tech Stack
    doc.add_heading('7. Technology Stack', level=1)
    ts = ['Groq / OpenRouter (LLM Inference)', 'Google Gemini (Text Embeddings)', 'Supabase pgvector (Vector DB)', 'FastAPI (Python Backend)', 'Render (Backend Hosting)', 'Next.js 14 App Router (React Frontend)', 'Vercel (Frontend Hosting)', 'Tailwind CSS & shadcn/ui (Styling)']
    for t in ts:
        doc.add_paragraph(t, style='List Bullet')

    # 8. Hard Rules
    doc.add_heading('8. Hard Rules', level=1)
    doc.add_paragraph('1. Never break character. Never acknowledge being an AI.')
    doc.add_paragraph('2. Token budgets must be strictly enforced via tiktoken to prevent context length errors on Groq.')
    doc.add_paragraph('3. All frontend styling must adhere to pure black (#0a0a0a) and crimson (#8B0000).')

    # 9. Deployment
    doc.add_heading('9. Deployment Checklist', level=1)
    doc.add_paragraph('Render (Backend): Create Web Service, attach GitHub repo, use `uvicorn app.main:app --host 0.0.0.0 --port 10000`. Ensure FRONTEND_URL is set for CORS.')
    doc.add_paragraph('Vercel (Frontend): Import repo, set Framework Preset to Next.js. Add NEXT_PUBLIC_API_URL pointing to the Render backend URL.')

    # 10. Known Issues
    doc.add_heading('10. Known Issues & Lessons Learned', level=1)
    doc.add_paragraph('Gemini 429 Quota: The free tier is limited to 1,000 embedding requests per day. The massive 4,400+ chunk ingestion exhausted this, causing chat to temporarily fail until the daily reset.')
    doc.add_paragraph('Tab Character Bug: The initial ingestion failed because raw tab characters in the Letters to Felice text file caused the chunker logic to halt and emit 0 bytes.')
    doc.add_paragraph('CI Working-Directory: GitHub Actions pytest workflow originally failed because it was executing in the root folder rather than the `backend/` directory.')

    doc.save(output_path)

# ── Document 2: Explainer ───────────────────────────────────────────────────
def create_explainer_doc(output_path):
    doc = Document()
    apply_styles(doc)

    title = doc.add_heading('ABDUL’S PROJECT EXPLAINER', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\nA clear, plain-language guide to pitching and explaining the Digital Clone AI.\n').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Pitch
    doc.add_heading('1. The Pitch', level=1)
    doc.add_paragraph().add_run('One-Line Pitch:').bold = True
    doc.add_paragraph('"I built a production-ready, full-stack AI platform that digitally resurrects historical figures using a multi-agent backend and a highly optimized memory retrieval system."')
    
    doc.add_paragraph().add_run('Full Paragraph Pitch:').bold = True
    doc.add_paragraph('This project is a multi-tenant Digital Clone AI architecture. It consists of a FastAPI Python backend integrated with a Supabase vector database, and a Next.js frontend. I engineered a strict 3-tier memory system and a robust context budgeting algorithm so the AI acts exactly like the target persona—in this case, Franz Kafka. By using Retrieval-Augmented Generation (RAG), the bot doesn’t just hallucinate; it accurately references over 4,000 real chunks of Kafka’s diaries and books in real time.')

    # 2. Services Explained
    doc.add_heading('2. Services Breakdown', level=1)
    services = [
        ('Groq / OpenRouter', 'The "brain". These are incredibly fast cloud providers that run the Llama 3 AI model to generate the text.'),
        ('Gemini API', 'The "translator". It converts human text into math (embeddings) so the database can understand it.'),
        ('Supabase pgvector', 'The "bookshelf". A PostgreSQL database in the cloud that stores all of Kafka\'s books and memories.'),
        ('FastAPI', 'The "manager". The Python backend framework that handles all the logic, budgeting, and database routing.'),
        ('Render', 'The "server room". Where the backend lives on the internet.'),
        ('Next.js / Vercel', 'The "face". Next.js runs the frontend user interface, and Vercel hosts it on the internet.'),
        ('Tailwind / shadcn', 'The "paint and tools". Used to make the UI look beautiful and dark.'),
        ('GitHub Actions', 'The "inspector". Automatically runs tests every time code is pushed to make sure nothing is broken.'),
        ('tiktoken', 'The "accountant". Counts the exact number of words (tokens) to ensure the AI doesn\'t crash from reading too much at once.')
    ]
    for s_name, s_desc in services:
        doc.add_paragraph().add_run(f'{s_name}: ').bold = True
        doc.paragraphs[-1].add_run(s_desc)

    # 3. RAG
    doc.add_heading('3. RAG (Retrieval-Augmented Generation) Explained', level=1)
    doc.add_paragraph('RAG is a way to make AI smart about specific topics without having to retrain the whole model. Here is how it works in our project:')
    steps = [
        'User types a message ("What is your relationship with your father?").',
        'Backend uses Gemini to turn that message into a list of numbers (an embedding).',
        'Backend sends those numbers to Supabase to find the closest matching numbers in our database.',
        'Supabase returns actual paragraphs from Kafka’s "Letter to His Father".',
        'Backend packages the user’s message AND those book paragraphs into one big prompt.',
        'Groq reads everything and responds intelligently in character.'
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'Step {i}: {s}')

    # 4. 3-Tier Memory Analogies
    doc.add_heading('4. 3-Tier Memory Analogies', level=1)
    doc.add_paragraph().add_run('Tier 1: Core Identity.').bold = True
    doc.paragraphs[-1].add_run(' This is the persona’s "DNA" or "Personality". It’s who they are at their core.')
    doc.add_paragraph().add_run('Tier 2: Episodic Memory (RAG).').bold = True
    doc.paragraphs[-1].add_run(' This is their "Bookshelf" or "Long-term memories". They can reach in and pull out specific quotes or events from the past.')
    doc.add_paragraph().add_run('Tier 3: Short-term History.').bold = True
    doc.paragraphs[-1].add_run(' This is the "Current Conversation". It’s remembering what was just said five minutes ago.')

    # 5. Cost
    doc.add_heading('5. Cost Breakdown ($0/month)', level=1)
    doc.add_paragraph('Supabase: Free tier (up to 500MB DB space).')
    doc.add_paragraph('Render / Vercel: Hobby free tiers for hosting.')
    doc.add_paragraph('Groq / Gemini: Free API tiers (which is why we hit rate limits if we upload too much too fast!).')

    # 6. Audience Scripts
    doc.add_heading('6. How to Explain This Project', level=1)
    doc.add_paragraph().add_run('To a non-technical person:').bold = True
    doc.add_paragraph('"I built a website where you can text Franz Kafka. Instead of a generic robot, I fed all of his private diaries and books into a custom database, so when he replies to you, he’s actually referencing his real memories and speaking in his exact historical voice."')
    doc.add_paragraph().add_run('To a technical interviewer:').bold = True
    doc.add_paragraph('"I engineered a multi-tenant RAG architecture using FastAPI and Next.js. I implemented strict token budgeting using tiktoken, semantic vector search via Supabase pgvector, and dynamically injected few-shot examples to enforce persona boundaries using Llama 3 via Groq."')
    doc.add_paragraph().add_run('To a developer friend:').bold = True
    doc.add_paragraph('"I built an AI clone backend. It uses Gemini for embeddings and pgvector for RAG. It handles strict context-window budgeting before streaming the response to a custom shadcn React frontend."')

    # 7. Quick Reference
    doc.add_heading('7. Quick Reference Commands', level=1)
    doc.add_paragraph('Backend Start: `uvicorn app.main:app --reload`')
    doc.add_paragraph('Frontend Start: `npm run dev`')
    doc.add_paragraph('Run Tests: `pytest`')
    doc.add_paragraph('Format Code: `ruff check . --fix`')

    doc.save(output_path)

if __name__ == '__main__':
    handoff_path = '/home/alucard/Downloads/AI Bot/Antigravity_Project_Handoff.docx'
    explainer_path = '/home/alucard/Downloads/AI Bot/Abdul_Project_Explainer.docx'
    
    create_handoff_doc(handoff_path)
    create_explainer_doc(explainer_path)
    print(f"Successfully generated {handoff_path} and {explainer_path}")
